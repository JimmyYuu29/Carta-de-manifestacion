import streamlit as st
import docx
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
import re
from datetime import datetime
import os
from copy import deepcopy
import io
import pandas as pd
from openpyxl import load_workbook

OFICINAS = {
    "ALICANTE": {
        "Direccion_Oficina": "Pintor Cabrera 22, esc. B, planta 4 A",
        "CP": "03003",
        "Ciudad_Oficina": "Alicante"
    },
    "BARCELONA": {
        "Direccion_Oficina": "C/ Diputació, 260",
        "CP": "08007",
        "Ciudad_Oficina": "Barcelona"
    },
    "BILBAO": {
        "Direccion_Oficina": "C/ Rodríguez Arias, 23, planta 6ª, Dpto. 12",
        "CP": "48011",
        "Ciudad_Oficina": "Bilbao"
    },
    "MADRID (Alcalá 63)": {
        "Direccion_Oficina": "C/ Alcalá, 63",
        "CP": "28014",
        "Ciudad_Oficina": "Madrid"
    },
    "MADRID (Alcalá 61-3ª)": {
        "Direccion_Oficina": "C/ Alcalá, 61, Planta 3ª",
        "CP": "28014",
        "Ciudad_Oficina": "Madrid"
    },
    "MÁLAGA": {
        "Direccion_Oficina": "Pirandello nº 6 portal 3, planta 6ª, puerta 4ª",
        "CP": "29010",
        "Ciudad_Oficina": "Málaga"
    },
    "VALENCIA": {
        "Direccion_Oficina": "C/ Félix Pizcueta, 4 – 4º",
        "CP": "46004",
        "Ciudad_Oficina": "Valencia"
    },
    "VIGO": {
        "Direccion_Oficina": "C/ República Argentina, 25 – 1º Izda",
        "CP": "36201",
        "Ciudad_Oficina": "Vigo"
    },
    "PERSONALIZADA": {           # Opción para escritura libre
        "Direccion_Oficina": "",
        "CP": "",
        "Ciudad_Oficina": ""
    }
}

# Configuración de la página
st.set_page_config(
    page_title="Generador de Cartas de Manifestación",
    page_icon="📄",
    layout="wide"
)

# Título principal
st.title("🏢 Generador de Cartas de Manifestación - Forvis Mazars")
st.markdown("---")

# Clase para manejar la generación de cartas
class CartaManifestacionGenerator:
    def __init__(self, template_path):
        self.template_path = template_path
        self.doc = Document(template_path)
        self.variables = {}
        self.conditionals = {}
        
    def extract_variables(self):
        """Extrae todas las variables y condicionales del documento"""
        variables = set()
        conditionals = set()
        
        # Buscar en párrafos
        for paragraph in self.doc.paragraphs:
            text = paragraph.text
            
            # Buscar variables {{variable}}
            var_matches = re.findall(r'\{\{([^}]+)\}\}', text)
            for match in var_matches:
                # Limpiar variable
                var_name = match.strip()
                # Caso especial para lista_alto_directores
                if 'lista_alto_directores' in var_name and ':' in var_name:
                    variables.add('lista_alto_directores')
                elif '|' in var_name:  # Manejar filtros como |int
                    var_name = var_name.split('|')[0].strip()
                    if not var_name.startswith('%'):
                        variables.add(var_name)
                elif not var_name.startswith('%'):  # Excluir código Jinja
                    variables.add(var_name)
            
            # Buscar condicionales {% if variable == 'valor' %}
            cond_matches = re.findall(r'\{%\s*if\s+(\w+)\s*==', text)
            for match in cond_matches:
                conditionals.add(match)
        
        # Buscar en tablas
        for table in self.doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    text = cell.text
                    var_matches = re.findall(r'\{\{([^}]+)\}\}', text)
                    for match in var_matches:
                        var_name = match.strip()
                        if 'lista_alto_directores' in var_name and ':' in var_name:
                            variables.add('lista_alto_directores')
                        elif '|' in var_name:
                            var_name = var_name.split('|')[0].strip()
                            if not var_name.startswith('%'):
                                variables.add(var_name)
                        elif not var_name.startswith('%'):
                            variables.add(var_name)
                    
                    cond_matches = re.findall(r'\{%\s*if\s+(\w+)\s*==', text)
                    for match in cond_matches:
                        conditionals.add(match)
        
        return sorted(list(variables)), sorted(list(conditionals))
    

    def _strip_conditional_blocks(self, doc, cond_values):
        """
        Borra todo el contenido (párrafos, tablas, etc.) comprendido entre
        {% if VAR == 'sí' %} ... {% endif %}
        cuando cond_values[VAR] == 'no'.
        Si cond_values[VAR] == 'sí', conserva el contenido pero elimina las líneas
        marcador {% if ... %} y {% endif %}.
        """
        body_elems = list(doc.element.body)     # Secuencia XML de párrafos y tablas
        inside_remove = False                   # Estamos dentro de un bloque a eliminar
        inside_keep   = False                   # Estamos dentro de un bloque a conservar
        trash = []                              # Elementos a suprimir del cuerpo

        for el in body_elems:
            # Obtener texto plano si es párrafo
            txt = ""
            if el.tag.endswith('p'):
                txt = "".join(t.text or "" for t in el.iter() if getattr(t, "text", None)).strip()

            # Apertura del bloque
            m_open = re.match(r"\{% if (\w+)\s*==\s*'sí' %\}", txt)
            if m_open:
                var = m_open.group(1)
                if cond_values.get(var, 'no') == 'sí':
                    inside_keep = True
                else:
                    inside_remove = True
                trash.append(el)               # Borra la línea marcador
                continue

            # Cierre del bloque
            if re.match(r"\{% endif %\}", txt):
                trash.append(el)               # Borra la línea marcador
                inside_remove = False
                inside_keep   = False
                continue

            # Elementos internos
            if inside_remove:
                trash.append(el)
            # inside_keep → se dejan intactos

        # Ejectutar removals
        for el in trash:
            el.getparent().remove(el)


    def process_template(self, variables, conditionals):
        """Procesa la plantilla con las variables proporcionadas"""
        new_doc = Document(self.template_path)
        
        self._strip_conditional_blocks(new_doc, conditionals)

        # Procesar párrafos
        for i, paragraph in enumerate(new_doc.paragraphs):
            original_text = paragraph.text
            if original_text.strip():
                new_text = self._replace_variables(original_text, variables, conditionals)
                if new_text != original_text:
                    # Guardar formato original
                    original_format = self._save_paragraph_format(paragraph)
                    
                    # Limpiar y aplicar nuevo texto
                    paragraph.clear()
                    paragraph.text = new_text
                    
                    # Restaurar formato
                    self._restore_paragraph_format(paragraph, original_format)
        
        # Procesar tablas
        for table in new_doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for paragraph in cell.paragraphs:
                        original_text = paragraph.text
                        if original_text.strip():
                            new_text = self._replace_variables(original_text, variables, conditionals)
                            if new_text != original_text:
                                paragraph.text = new_text
        
        # Corregir numeración
        self._fix_numbering(new_doc)
        self._remove_underlines(new_doc)  
        return new_doc
    
    def _replace_variables(self, text, variables, conditionals):
        """Reemplaza variables y procesa condicionales"""
        # Primero procesar condicionales
        text = self._process_conditionals(text, conditionals)
        
        # Manejar caso especial de lista_alto_directores con formato complejo
        # Buscar el patrón completo que incluye todo el bloque con ejemplos
        lista_pattern = r'\{\{lista_alto_directores:[^}]+\}\}'
        lista_matches = list(re.finditer(lista_pattern, text, re.DOTALL))
        
        # Reemplazar de atrás hacia adelante para no afectar los índices
        for match in reversed(lista_matches):
            if 'lista_alto_directores' in variables and variables['lista_alto_directores']:
                text = text[:match.start()] + variables['lista_alto_directores'] + text[match.end():]
            else:
                text = text[:match.start()] + text[match.end():]
        
        # Luego reemplazar variables simples
        for var_name, var_value in variables.items():
            if var_name == 'lista_alto_directores':
                continue  # Ya procesada arriba
                
            # Manejar diferentes patrones de variables
            patterns = [
                rf'\{{\{{\s*{re.escape(var_name)}\s*\}}\}}',
                rf'\{{\{{\s*{re.escape(var_name)}\s*\|\s*int\s*\}}\}}',
                rf'\{{\{{\s*{re.escape(var_name)}\s*\|\s*int\s*-\s*1\s*\}}\}}'
            ]
            
            for pattern in patterns:
                if '|int - 1' in pattern and var_value:
                    try:
                        replacement = str(int(var_value) - 1)
                    except:
                        replacement = var_value
                else:
                    replacement = str(var_value) if var_value else ''
                
                text = re.sub(pattern, replacement, text)
        
        # Limpiar marcadores restantes
        text = re.sub(r'\[?\{\{[^}]*\}\}\]?', '', text)
        text = re.sub(r'\[\]\.mark', '', text)
        text = re.sub(r'\.mark', '', text)
        text = re.sub(r'\[\.mark\]', '', text)
        
        return text
    
    def _process_conditionals(self, text, conditionals):
        """Procesa bloques condicionales"""
        for cond_var, cond_value in conditionals.items():
            # Patrón para bloques if con mark
            if_pattern = rf'\[\{{% if {cond_var} == \'sí\' %\}}\]\.mark(.*?)\[\{{% endif %\}}\]\.mark'
            if cond_value == 'sí':
                text = re.sub(if_pattern, r'\1', text, flags=re.DOTALL)
            else:
                text = re.sub(if_pattern, '', text, flags=re.DOTALL)
            
            # Patrón para bloques if sin mark
            if_pattern = rf'\{{% if {cond_var} == \'sí\' %\}}(.*?)\{{% endif %\}}'
            if cond_value == 'sí':
                text = re.sub(if_pattern, r'\1', text, flags=re.DOTALL)
            else:
                text = re.sub(if_pattern, '', text, flags=re.DOTALL)
        
        # Limpiar marcas de condicionales restantes
        text = re.sub(r'\{%[^%]*%\}', '', text)
        
        return text
    
    def _save_paragraph_format(self, paragraph):
        """Guarda el formato de un párrafo"""
        format_info = {
            'alignment': paragraph.alignment,
            'style': paragraph.style.name if paragraph.style else None,
            'runs': []
        }
        
        for run in paragraph.runs:
            run_format = {
                'bold': run.bold,
                'italic': run.italic,
                'underline': run.underline,
                'font_name': run.font.name,
                'font_size': run.font.size,
                'font_color': run.font.color.rgb if run.font.color and run.font.color.rgb else None
            }
            format_info['runs'].append(run_format)
        
        return format_info
    
    def _restore_paragraph_format(self, paragraph, format_info):
        """Restaura el formato de un párrafo"""
        if format_info['alignment']:
            paragraph.alignment = format_info['alignment']
        
        if format_info['style']:
            try:
                paragraph.style = format_info['style']
            except:
                pass
        
        # Si hay información de runs, intentar aplicar formato
        if format_info['runs'] and paragraph.runs:
            for i, run in enumerate(paragraph.runs):
                if i < len(format_info['runs']):
                    run_format = format_info['runs'][i]
                    if run_format['bold'] is not None:
                        run.bold = run_format['bold']
                    if run_format['italic'] is not None:
                        run.italic = run_format['italic']
                    if run_format['underline'] is not None:
                        run.underline = run_format['underline']
                    if run_format['font_name']:
                        run.font.name = run_format['font_name']
                    if run_format['font_size']:
                        run.font.size = run_format['font_size']
    
    def _fix_numbering(self, doc):
        """Corrige la numeración de los puntos"""
        current_number = 1
        sub_number = 1
        in_sub_list = False
        
        for paragraph in doc.paragraphs:
            text = paragraph.text.strip()
            
            # Detectar puntos principales (empiezan con número seguido de punto)
            main_match = re.match(r'^(\d+)\.\s+(.+)', text)
            if main_match:
                paragraph.text = f"{current_number}. {main_match.group(2)}"
                current_number += 1
                in_sub_list = False
            
            # Detectar sub-puntos (empiezan con letra minúscula seguida de punto)
            sub_match = re.match(r'^[a-z]\.\s+(.+)', text)
            if sub_match:
                if not in_sub_list:
                    sub_number = 1
                    in_sub_list = True
                
                letter = chr(ord('a') + sub_number - 1)
                paragraph.text = f"{letter}. {sub_match.group(1)}"
                sub_number += 1
    
    def _remove_underlines(self, doc):
        """Quita cualquier subrayado que se haya quedado en runs"""
        # Párrafos normales
        for p in doc.paragraphs:
            for run in p.runs:
                run.underline = False
        # Párrafos dentro de tablas
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for p in cell.paragraphs:
                        for run in p.runs:
                            run.underline = False

# Función que permite a subir ficheros Excel o Word y extraer variable
def process_uploaded_file(uploaded_file, file_type):
    """
    Procesa archivos Excel o Word subidos y extrae las variables y sus valores
    
    Returns:
        dict: Diccionario con las variables/condiciones y sus valores
    """
    extracted_data = {}
    
    try:
        if file_type == "excel":
            # Leer Excel con pandas
            df = pd.read_excel(uploaded_file, header=None)
            
            # Verificar que tenga al menos 2 columnas
            if df.shape[1] >= 2:
                # Iterar por las filas
                for index, row in df.iterrows():
                    if pd.notna(row[0]) and pd.notna(row[1]):
                        var_name = str(row[0]).strip()
                        var_value = row[1]
                        
                        # Si es una fecha/datetime, convertirla a string
                        if pd.api.types.is_datetime64_any_dtype(type(var_value)) or isinstance(var_value, datetime):
                            var_value = var_value.strftime("%d/%m/%Y")
                        else:
                            var_value = str(var_value).strip()

                        # Convertir 'SI'/'NO'/1/0 a 'sí'/'no' para condiciones
                        if var_value.upper() in ['SI', 'SÍ'] or var_value == '1':
                            var_value = 'sí'
                        elif var_value.upper() == 'NO' or var_value == '0':
                            var_value = 'no'
                        
                        # Normalizar nombres de variables
                        normalized_name = var_name.lower().replace('á', 'a').replace('é', 'e').replace('í', 'i').replace('ó', 'o').replace('ú', 'u')
                        
                        # Mapeo de nombres alternativos
                        name_mapping = {
                            'comision': 'comision',
                            'comisión': 'comision',
                            'organo': 'organo',
                            'órgano': 'organo'
                        }
                        
                        # Usar el nombre normalizado si existe en el mapeo
                        if normalized_name in name_mapping:
                            var_name = name_mapping[normalized_name]
                        
                        extracted_data[var_name] = var_value
        
        elif file_type == "word":
            # Leer documento Word
            doc = Document(uploaded_file)
            
            # Extraer texto de todos los párrafos
            for paragraph in doc.paragraphs:
                text = paragraph.text.strip()
                if text and ':' in text:
                    # Dividir por el primer ':' encontrado
                    parts = text.split(':', 1)
                    if len(parts) == 2:
                        var_name = parts[0].strip()
                        var_value = parts[1].strip()
                        
                        # Convertir 'SI'/'NO'/1/0 a 'sí'/'no' para condiciones
                        if var_value.upper() in ['SI', 'SÍ'] or var_value == '1':
                            var_value = 'sí'
                        elif var_value.upper() == 'NO' or var_value == '0':
                            var_value = 'no'
                        
                        # Normalizar nombres de variables
                        normalized_name = var_name.lower().replace('á', 'a').replace('é', 'e').replace('í', 'i').replace('ó', 'o').replace('ú', 'u')
                        
                        # Mapeo de nombres alternativos
                        name_mapping = {
                            'comision': 'comision',
                            'comisión': 'comision',
                            'organo': 'organo',
                            'órgano': 'organo'
                        }
                        
                        # Usar el nombre normalizado si existe en el mapeo
                        if normalized_name in name_mapping:
                            var_name = name_mapping[normalized_name]
                        
                        extracted_data[var_name] = var_value
    
    except Exception as e:
        st.error(f"Error al procesar el archivo: {str(e)}")
        return {}
    
    return extracted_data


def parse_date_string(date_string):
    """Convierte string de fecha a objeto datetime"""
    if not date_string:
        return datetime.now()
    
    # Intentar varios formatos de fecha
    formats = [
        "%d/%m/%Y",
        "%Y-%m-%d", 
        "%d-%m-%Y",
        "%d de %B de %Y",
        "%Y/%m/%d",
        "%d.%m.%Y",
        "%Y.%m.%d",
        "%m/%d/%Y",  # Formato americano
        "%Y/%d/%m",  # Año/día/mes
    ]
    
    # Para manejar nombres de meses en español
    import locale
    try:
        locale.setlocale(locale.LC_TIME, 'es_ES.UTF-8')
    except:
        try:
            locale.setlocale(locale.LC_TIME, 'Spanish_Spain')
        except:
            pass
    
    for fmt in formats:
        try:
            return datetime.strptime(date_string, fmt)
        except:
            continue
    
    # Si no se puede parsear, devolver fecha actual
    return datetime.now()


def main():
    # Verificar si existe el archivo de plantilla
    template_path = "Modelo de plantilla.docx"
    
    if not os.path.exists(template_path):
        st.error(f"⚠️ No se encontró el archivo '{template_path}' en la carpeta actual.")
        st.info("Por favor, asegúrate de que el archivo de plantilla esté en la misma carpeta que este script.")
        return
    
    # Crear generador
    generator = CartaManifestacionGenerator(template_path)
    
    # Extraer variables y condicionales
    with st.spinner("Analizando plantilla..."):
        variables, conditionals = generator.extract_variables()
    
    st.success(f"✅ Plantilla analizada. Se encontraron {len(variables)} variables y {len(conditionals)} condicionales.")
    
    # Crear formulario en columnas
    st.subheader("📝 Información de la Carta")
    
    # Diccionarios para almacenar valores
    var_values = {}
    cond_values = {}
    
    # SECCIÓN DE IMPORTACIÓN DE DATOS
    st.markdown("---")
    st.subheader("📁 Importar datos desde archivo")
    
    col_import1, col_import2 = st.columns(2)
    
    with col_import1:
        uploaded_excel = st.file_uploader(
            "Cargar archivo Excel (.xlsx, .xls)", 
            type=['xlsx', 'xls'],
            help="Formato: Columna 1 = Nombre variable, Columna 2 = Valor"
        )
    
    with col_import2:
        uploaded_word = st.file_uploader(
            "Cargar archivo Word (.docx)", 
            type=['docx'],
            help="Formato: nombre_variable: valor (una por línea)"
        )
    
    # Procesar archivos si se han cargado
    imported_data = {}
    
    if uploaded_excel is not None:
        with st.spinner("Procesando archivo Excel..."):
            imported_data = process_uploaded_file(uploaded_excel, "excel")
            if imported_data:
                st.success(f"✅ Se importaron {len(imported_data)} valores desde Excel")
    
    elif uploaded_word is not None:
        with st.spinner("Procesando archivo Word..."):
            imported_data = process_uploaded_file(uploaded_word, "word")
            if imported_data:
                st.success(f"✅ Se importaron {len(imported_data)} valores desde Word")
    

    # Aplicar datos importados a las variables y condiciones
    if imported_data:
        # Separar variables y condiciones basándose en las listas extraídas
        for key, value in imported_data.items():
            if key in conditionals:
                cond_values[key] = value
            else:
                var_values[key] = value
    
    st.markdown("---")
    # FIN DE LA SECCIÓN DE IMPORTACIÓN
    
    # Organizar variables por categorías
    col1, col2 = st.columns(2)
    
    with col1:
        st.markdown("### 📋 Información de la Oficina")
        oficina_sel = st.selectbox(
            "Selecciona la oficina",
            list(OFICINAS.keys()),
            index=list(OFICINAS.keys()).index(
                var_values.get("Oficina_Seleccionada", "BARCELONA")
            )
        )
        var_values["Oficina_Seleccionada"] = oficina_sel
           
        sel_data = OFICINAS[oficina_sel]
        for campo in ["Direccion_Oficina", "CP", "Ciudad_Oficina"]:
            # Sólo pisamos valores si el usuario no había importado uno distinto
            if campo not in var_values or var_values[campo] == "":
                var_values[campo] = sel_data[campo]
    
        # --- Mostrar / permitir edición -----------------------------------------
        edicion_libre = (oficina_sel == "PERSONALIZADA")
        var_values["Direccion_Oficina"] = st.text_input(
            "Dirección de la Oficina",
            value=var_values.get("Direccion_Oficina", ""),
            key="direccion",
            disabled=not edicion_libre
        )
        var_values["CP"] = st.text_input(
            "Código Postal",
            value=var_values.get("CP", ""),
            key="cp",
            disabled=not edicion_libre
        )
        var_values["Ciudad_Oficina"] = st.text_input(
            "Ciudad",
            value=var_values.get("Ciudad_Oficina", ""),
            key="ciudad",
            disabled=not edicion_libre
        )

        st.markdown("### 🏢 Nombre de cliente")
        var_values['Nombre_Cliente'] = st.text_input(
            "Nombre del Cliente", 
            value=var_values.get('Nombre_Cliente', ''),
            key="nombre_cliente"
        )

        st.markdown("### 📅 Fechas")
        # Fecha de hoy
        fecha_hoy = parse_date_string(var_values.get('Fecha_de_hoy', ''))
        var_values['Fecha_de_hoy'] = st.date_input(
            "Fecha de Hoy", 
            value=fecha_hoy
        ).strftime("%d de %B de %Y")
        
        # Fecha del encargo
        fecha_encargo = parse_date_string(var_values.get('Fecha_encargo', ''))
        var_values['Fecha_encargo'] = st.date_input(
            "Fecha del Encargo",
            value=fecha_encargo
        ).strftime("%d de %B de %Y")
        
        # Fecha fin del ejercicio
        fecha_ff = parse_date_string(var_values.get('FF_Ejecicio', ''))
        var_values['FF_Ejecicio'] = st.date_input(
            "Fecha Fin del Ejercicio",
            value=fecha_ff
        ).strftime("%d de %B de %Y")
        
        # Fecha de cierre
        fecha_cierre = parse_date_string(var_values.get('Fecha_cierre', ''))
        var_values['Fecha_cierre'] = st.date_input(
            "Fecha de Cierre",
            value=fecha_cierre
        ).strftime("%d de %B de %Y")
        
        st.markdown("### 📝 Información General")
        var_values['Lista_Abogados'] = st.text_area(
            "Lista de abogados y asesores fiscales", 
            value=var_values.get('Lista_Abogados', ''),
            placeholder="Ej: Despacho ABC - Asesoría fiscal\nDespacho XYZ - Asesoría legal",
            key="abogados"
        )
        var_values['anexo_partes'] = st.text_input(
            "Número anexo partes vinculadas", 
            value=var_values.get('anexo_partes', '2')
        )
        var_values['anexo_proyecciones'] = st.text_input(
            "Número anexo proyecciones", 
            value=var_values.get('anexo_proyecciones', '3')
        )

    with col2:
        st.markdown("### 👥 Órgano de Administración")
        # Usar el valor importado si existe, sino usar 'consejo' por defecto
        if 'organo' in imported_data:
            organo_default = imported_data['organo']
        else:
            organo_default = cond_values.get('organo', 'consejo')
        cond_values['organo'] = st.selectbox(
            "Tipo de Órgano de Administración",
            options=['consejo', 'administrador_unico', 'administradores'],
            index=['consejo', 'administrador_unico', 'administradores'].index(organo_default),
            format_func=lambda x: {
                'consejo': 'Consejo de Administración',
                'administrador_unico': 'Administrador Único',
                'administradores': 'Administradores'
            }[x]
        )
        
        st.markdown("### ✅ Opciones Condicionales")
        
        # Comisión de Auditoría
        cond_values['comision'] = 'sí' if st.checkbox(
            "¿Existe Comisión de Auditoría?",
            value=(cond_values.get('comision', 'no') == 'sí')
        ) else 'no'
        
        # Junta y Comité
        cond_values['junta'] = 'sí' if st.checkbox(
            "¿Incluir Junta de Accionistas?",
            value=(cond_values.get('junta', 'no') == 'sí')
        ) else 'no'
        
        cond_values['comite'] = 'sí' if st.checkbox(
            "¿Incluir Comité?",
            value=(cond_values.get('comite', 'no') == 'sí')
        ) else 'no'
        
        # Incorrecciones
        cond_values['incorreccion'] = 'sí' if st.checkbox(
            "¿Hay incorrecciones no corregidas?",
            value=(cond_values.get('incorreccion', 'no') == 'sí')
        ) else 'no'
        
        if cond_values['incorreccion'] == 'sí':
            with st.container():
                st.markdown("##### 📌 Detalles de incorrecciones")
                var_values['Anio_incorreccion'] = st.text_input(
                    "Año de la incorrección", 
                    value=var_values.get('Anio_incorreccion', ''),
                    key="anio_inc"
                )
                var_values['Epigrafe'] = st.text_input(
                    "Epígrafe afectado", 
                    value=var_values.get('Epigrafe', ''),
                    key="epigrafe"
                )
                if 'limitacion_alcance' not in cond_values:
                    cond_values['limitacion_alcance'] = 'no'
                cond_values['limitacion_alcance'] = 'sí' if st.checkbox(
                    "¿Hay limitación al alcance?",
                    value=(cond_values.get('limitacion_alcance', 'no') == 'sí')
                ) else 'no'
                if cond_values['limitacion_alcance'] == 'sí':
                    var_values['detalle_limitacion'] = st.text_area(
                        "Detalle de la limitación", 
                        value=var_values.get('detalle_limitacion', ''),
                        key="det_limitacion"
                    )
        
        # Dudas empresa en funcionamiento
        cond_values['dudas'] = 'sí' if st.checkbox(
            "¿Existen dudas sobre empresa en funcionamiento?",
            value=(cond_values.get('dudas', 'no') == 'sí')
        ) else 'no'
        
        # Arrendamientos
        cond_values['rent'] = 'sí' if st.checkbox(
            "¿Incluir párrafo sobre arrendamientos?",
            value=(cond_values.get('rent', 'no') == 'sí')
        ) else 'no'

        # Valor razonable a Coste
        cond_values['A_coste'] = 'sí' if st.checkbox(
            "¿Hay activos valorados a coste en vez de valor razonable?",
            value=(cond_values.get('A_coste', 'no') == 'sí')
        ) else 'no'
        
        # Experto independiente
        cond_values['experto'] = 'sí' if st.checkbox(
            "¿Se utilizó un experto independiente?",
            value=(cond_values.get('experto', 'no') == 'sí')
        ) else 'no'
        
        if cond_values['experto'] == 'sí':
            with st.container():
                st.markdown("##### 📌 Información del experto")
                var_values['nombre_experto'] = st.text_input(
                    "Nombre del experto", 
                    value=var_values.get('nombre_experto', ''),
                    key="experto_nombre"
                )
                var_values['experto_valoracion'] = st.text_input(
                    "Elemento valorado por experto", 
                    value=var_values.get('experto_valoracion', ''),
                    key="experto_val"
                )
        
        # Unidad de decisión
        cond_values['unidad_decision'] = 'sí' if st.checkbox(
            "¿Bajo la misma unidad de decisión?",
            value=(cond_values.get('unidad_decision', 'no') == 'sí')
        ) else 'no'
        
        if cond_values['unidad_decision'] == 'sí':
            with st.container():
                st.markdown("##### 📌 Información de la unidad de decisión")
                var_values['nombre_unidad'] = st.text_input(
                    "Nombre de la unidad", 
                    value=var_values.get('nombre_unidad', ''),
                    key="nombre_unidad"
                )
                var_values['nombre_mayor_sociedad'] = st.text_input(
                    "Nombre de la mayor sociedad", 
                    value=var_values.get('nombre_mayor_sociedad', ''),
                    key="nombre_mayor_sociedad"
                )
                var_values['localizacion_mer'] = st.text_input(
                    "Localización o domiciliación mercantil", 
                    value=var_values.get('localizacion_mer', ''),
                    key="localizacion_mer"
                )

        # Activos por impuestos
        cond_values['activo_impuesto'] = 'sí' if st.checkbox(
            "¿Hay activos por impuestos diferidos?",
            value=(cond_values.get('activo_impuesto', 'no') == 'sí')
        ) else 'no'
        
        if cond_values['activo_impuesto'] == 'sí':
            with st.container():
                st.markdown("##### 📌 Recuperación de activos")
                var_values['ejercicio_recuperacion_inicio'] = st.text_input(
                    "Ejercicio inicio recuperación", 
                    value=var_values.get('ejercicio_recuperacion_inicio', ''),
                    key="rec_inicio"
                )
                var_values['ejercicio_recuperacion_fin'] = st.text_input(
                    "Ejercicio fin recuperación", 
                    value=var_values.get('ejercicio_recuperacion_fin', ''),
                    key="rec_fin"
                )
        
        # Operaciones en paraísos fiscales
        cond_values['operacion_fiscal'] = 'sí' if st.checkbox(
            "¿Operaciones en paraísos fiscales?",
            value=(cond_values.get('operacion_fiscal', 'no') == 'sí')
        ) else 'no'
        
        if cond_values['operacion_fiscal'] == 'sí':
            with st.container():
                st.markdown("##### 📌 Detalle operaciones")
                var_values['detalle_operacion_fiscal'] = st.text_area(
                    "Detalle operaciones paraísos fiscales", 
                    value=var_values.get('detalle_operacion_fiscal', ''),
                    key="det_fiscal"
                )
        
        # Compromisos por pensiones
        cond_values['compromiso'] = 'sí' if st.checkbox(
            "¿Compromisos por pensiones?",
            value=(cond_values.get('compromiso', 'no') == 'sí')
        ) else 'no'
        
        # Informe de gestión
        cond_values['gestion'] = 'sí' if st.checkbox(
            "¿Incluir informe de gestión?",
            value=(cond_values.get('gestion', 'no') == 'sí')
        ) else 'no'

    # Lista de altos directivos - Sección separada
    st.markdown("---")
    st.markdown("### 👔 Alta Dirección")
    
    # Verificar si se importó lista_alto_directores
    imported_directivos = var_values.get('lista_alto_directores', '')
    
    if imported_directivos:
        # Parsear la lista importada
        directivos_lines = imported_directivos.strip().split('\n')
        directivos_list = []
        for line in directivos_lines:
            if line.strip():
                directivos_list.append(line.strip())
        
        st.info(f"📌 Se importaron {len(directivos_list)} directivos")
        st.text_area("Directivos importados:", value='\n'.join(directivos_list), height=100, disabled=True)
        
        # Preguntar si quiere usar los importados o editarlos
        use_imported = st.checkbox("Usar directivos importados", value=True)
        
        if not use_imported:
            # Permitir edición manual
            st.info("Introduce los nombres y cargos de los altos directivos.")
            num_directivos = st.number_input("Número de altos directivos", min_value=0, max_value=10, value=len(directivos_list))
            
            directivos_list = []
            indent = "                                  "
            for i in range(num_directivos):
                col_nombre, col_cargo = st.columns(2)
                with col_nombre:
                    nombre = st.text_input(f"Nombre completo {i+1}", key=f"dir_nombre_{i}")
                with col_cargo:
                    cargo = st.text_input(f"Cargo {i+1}", key=f"dir_cargo_{i}")
                if nombre and cargo:
                    directivos_list.append(f"{indent} D. {nombre} - {cargo}")
    else:
        # No hay directivos importados, entrada manual
        st.info("Introduce los nombres y cargos de los altos directivos. Estos reemplazarán completamente el ejemplo en la plantilla.")
        
        num_directivos = st.number_input("Número de altos directivos", min_value=0, max_value=10, value=2)
        
        directivos_list = []
        indent = "                                  "
        for i in range(num_directivos):
            col_nombre, col_cargo = st.columns(2)
            with col_nombre:
                nombre = st.text_input(f"Nombre completo {i+1}", key=f"dir_nombre_{i}")
            with col_cargo:
                cargo = st.text_input(f"Cargo {i+1}", key=f"dir_cargo_{i}")
            if nombre and cargo:
                directivos_list.append(f"{indent} D. {nombre} - {cargo}")
    
    st.markdown("---")
    st.markdown("### 👥 Persona de firma")
    var_values['Nombre_Firma'] = st.text_input(
        "Nombre del firmante", 
        value=var_values.get('Nombre_Firma', ''),
        key="nombre_firma"
    )
    var_values['Cargo_Firma'] = st.text_input(
        "Cargo del firmante", 
        value=var_values.get('Cargo_Firma', ''),
        key="cargo_firma"
    )

    # Mapa de dependencias
    conditional_vars_map = {
        'incorreccion': ['Anio_incorreccion', 'Epigrafe', 'detalle_limitacion'],
        'experto': ['nombre_experto', 'experto_valoracion'],
        'activo_impuesto': ['ejercicio_recuperacion_inicio', 'ejercicio_recuperacion_fin'],
        'operacion_fiscal': ['detalle_operacion_fiscal'],
        'unidad_decision': ['nombre_unidad', 'nombre_mayor_sociedad', 'localizacion_mer']
    }

    # Determinar variables requeridas
    required_vars = set(variables)
    for cond, vlist in conditional_vars_map.items():
        if cond_values.get(cond, 'no') == 'no':
            required_vars -= set(vlist)

    # Revisión de variables y condiciones
    st.markdown("---")
    st.header("🔍 Revisión automática")

    # Calcular los que faltan
    missing_vars = [v for v in required_vars if v not in var_values or var_values[v] == ""]
    missing_conds = [c for c in conditionals if c not in cond_values]

    # Mostrar resumen de importación si hubo datos importados
    if imported_data:
        st.info(f"📊 Datos importados: {len([k for k in imported_data if k in var_values])} variables y {len([k for k in imported_data if k in cond_values])} condiciones")

    # Informar al usuario
    if not missing_vars and not missing_conds:
        st.success("✅ Todas las variables y condiciones están completas.")
    else:
        st.warning(f"⚠️ Faltan {len(missing_vars)} variables y {len(missing_conds)} condiciones.")
        
        # Mostrar detalle y ofrecer campos para rellenar
        with st.expander("Ver / completar elementos faltantes"):
            # Inputs dinámicos para variables pendientes
            for var in missing_vars:
                var_values[var] = st.text_input(f"Valor para «{var}»", key=f"auto_{var}")

            # Checkboxes para condiciones pendientes
            for cond in missing_conds:
                cond_values[cond] = 'sí' if st.checkbox(f"Activar condición «{cond}»", key=f"auto_{cond}") else 'no'

    # Asignar la lista formateada
    var_values['lista_alto_directores'] = "\n".join(directivos_list) 
    
    # Vista previa de directivos
    if directivos_list:
        st.markdown("#### Vista previa de la lista de directivos:")
        st.code("\n".join(directivos_list))
    
    # Botón para generar carta
    st.markdown("---")
    
    if st.button("🚀 Generar Carta de Manifestación", type="primary"):
        # Validar campos obligatorios
        required_fields = ['Nombre_Cliente', 'Direccion_Oficina', 'CP', 'Ciudad_Oficina']
        missing_fields = [field for field in required_fields if not var_values.get(field)]
        
        if missing_fields:
            st.error(f"⚠️ Por favor completa los siguientes campos obligatorios: {', '.join(missing_fields)}")
        else:
            with st.spinner("Generando carta..."):
                try:
                    # Asegurar que todas las variables tengan un valor
                    for var in variables:
                        if var not in var_values:
                            var_values[var] = ""
                    
                    # Combinar variables regulares y condicionales
                    all_vars = {**var_values, **cond_values}
                    
                    # Procesar plantilla
                    new_doc = generator.process_template(all_vars, cond_values)
                    
                    # Guardar en memoria
                    doc_buffer = io.BytesIO()
                    new_doc.save(doc_buffer)
                    doc_buffer.seek(0)
                    
                    # Generar nombre de archivo
                    filename = f"Carta_Manifestacion_{var_values['Nombre_Cliente'].replace(' ', '_')}_{datetime.now().strftime('%Y%m%d')}.docx"
                    
                    st.success("✅ Carta generada exitosamente!")
                    
                    # Botón de descarga
                    st.download_button(
                        label="📥 Descargar Carta de Manifestación",
                        data=doc_buffer,
                        file_name=filename,
                        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                    )
                    
                except Exception as e:
                    st.error(f"❌ Error al generar la carta: {str(e)}")
                    st.exception(e)

# Ejecutar aplicación
if __name__ == "__main__":
    main()
